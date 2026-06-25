from pathlib import Path

from docx import Document


def _table_text(table) -> str:
    rows = []

    for row in table.rows:
        cells = [
            " ".join(cell.text.split())
            for cell in row.cells
            if cell.text and cell.text.strip()
        ]

        if cells:
            rows.append(" | ".join(cells))

    return "\n".join(rows)


def extract_docx_pages(file_path: Path) -> list[dict]:
    document = Document(file_path)

    text_blocks = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            text_blocks.append(text)

    for table in document.tables:
        text = _table_text(table).strip()

        if text:
            text_blocks.append(text)

    text = "\n\n".join(text_blocks).strip()

    if not text:
        return []

    return [{
        "page_number": 1,
        "text": text,
        "extraction_method": "text",
    }]
